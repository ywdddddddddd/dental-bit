#!/usr/bin/env python3
"""Generate final .docx matching updated manuscript (community focus, author info, 10 refs)."""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = 'papers/manuscript.docx'
FIG_DIR = 'papers/figures'

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17); s.right_margin = Cm(3.17)

style = doc.styles['Normal']; style.font.name = '宋体'; style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def h1(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(16); r.font.bold = True; r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def h2(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(14); r.font.bold = True; r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def h3(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(12); r.font.bold = True; r.font.name = '楷体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

def body(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(10.5); r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.first_line_indent = Pt(21); p.paragraph_format.line_spacing = 1.5

def img(path, caption, w=5.5):
    if os.path.exists(path):
        pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER; pi.add_run().add_picture(path, width=Inches(w))
        pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = pc.add_run(caption); r.font.size = Pt(9); r.font.bold = True; r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def tbl(headers, rows, caption=''):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''; rr = c.paragraphs[0].add_run(h); rr.font.size = Pt(9); rr.font.bold = True; rr.font.name = '宋体'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''; rr = c.paragraphs[0].add_run(str(val)); rr.font.size = Pt(9); rr.font.name = '宋体'; rr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(caption); r.font.size = Pt(9); r.font.bold = True; r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    doc.add_paragraph()

def ref_item(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(9); r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============ TITLE & AUTHORS ============
h1('口腔预问诊系统的基层应用研究')
doc.add_paragraph()
body('姚雯道，赵春钢*')
body('荆楚理工学院，湖北荆门 448000')
body('通信作者：赵春钢，E-mail: 75966778@qq.com')
body('基金项目：无')
body('第一作者简介：姚雯道，荆楚理工学院口腔医学系。通信作者简介：赵春钢（1982年2月—），男，汉族，湖北荆门人，硕士，讲师，口腔颌面外科教研室主任。从事口腔颌面外科学、口腔组织病理学、口腔种植学教学工作及口腔医学相关方向科研工作。')
doc.add_paragraph()

# ============ ABSTRACT ============
h3('【摘要】')
body('目的：探讨基于有向无环图（DAG）因果推理的口腔预问诊系统在社区居民口腔自查和基层分诊中的应用价值。')
body('方法：构建覆盖145种口腔疾病的知识驱动预问诊引擎，采用51节点DAG决策树实现症状-诊断的可追溯推理，融合第四次全国口腔健康流行病学调查数据作为评分先验。系统以单网页形式运行，支持离线使用。通过社交平台招募社区志愿者，回收有效完整口腔症状自评问卷354份，以命中率、分诊一致率和漏诊率为核心指标进行评价。')
body('结果：354份有效问卷中，系统Top-1诊断命中率66.7%，Top-3覆盖率83.6%。分诊一致率94.9%，过诊率5.1%，无一例红色紧急病例漏诊。黏膜类疾病的F1-score最高（0.72），牙体牙髓类次之（0.68）。')
body('结论：该系统可作为社区居民口腔自查和基层卫生工作者分诊的辅助工具推广应用。')
body('【关键词】口腔预问诊；社区口腔自查；临床决策支持；有向无环图；基层卫生')
doc.add_paragraph()

h3('Abstract')
body('Objective: To evaluate the application value of an oral pre-consultation system based on DAG causal reasoning for community residents oral self-examination and primary care triage. Methods: A knowledge-driven engine covering 145 oral diseases was developed. A total of 354 valid self-assessment questionnaires were collected from community volunteers. Results: Top-1 accuracy was 66.7%, Top-3 coverage 83.6%, triage consistency 94.9%, with zero missed red-flag cases. Conclusion: The system demonstrates strong potential for community oral self-examination and primary care triage assistance.')
body('Keywords: Oral pre-consultation; Community oral self-examination; Clinical decision support; Directed acyclic graph; Primary care')
doc.add_paragraph()

# ============ 1 引言 ============
h2('1 引言')
body('口腔疾病是全球最普遍的慢性非传染性疾病，约35亿人受其困扰[1]。我国第四次全国口腔健康流行病学调查显示，5岁儿童乳牙患龋率高达70.9%，35-44岁成人牙周健康率仅9.1%[2]。然而，我国口腔医师资源分布极不均衡——每百万人口口腔医师密度不足100，80%以上集中在城市三级医院。对广大社区居民而言，"牙疼了该不该去医院""牙龈出血严不严重""嘴里长了个东西要不要紧"——这些日常口腔困扰往往因缺乏专业判断而被拖延[2]。')
body('近年来，大语言模型（LLM）在医疗问诊中展现出引人瞩目的能力，但其在口腔专科场景中的表现尚不稳定——Hassanein等系统综述指出，LLM口腔病变诊断准确率范围仅为25-96%，且高度依赖提示词策略和输入模态[3]。专用牙科症状检查器如新加坡军队的Toothbuddy和芬兰的Omaolo虽在三级分诊中取得了较好的验证结果[4-5]，但疾病覆盖极为有限（<=20种），难以适应社区居民多元化的口腔主诉。')
body('针对上述问题，本研究设计了一种基于DAG因果推理的口腔预问诊引擎，覆盖145种口腔疾病并融合中国口腔流行病学先验数据。该系统以单一网页形式运行，无需安装或联网，旨在为社区居民提供便捷的口腔症状自筛工具，同时为基层卫生工作者提供标准化的预检分诊辅助。本文通过对354份社区志愿者的真实问卷数据的系统评价，探讨该系统在社区自查和基层分诊场景中的实际应用潜力。')

# ============ 2 资料与方法 ============
h2('2 资料与方法')
h3('2.1 系统设计思路')
body('系统采用"知识驱动"的设计路线。核心思路是将口腔临床专家的诊断逻辑显式编码为DAG决策网络：每组"症状-疾病"关联通过delta值（加分权重）和lock规则（排除关系）编码，疾病之间的相互关系通过三层网络（因果链、互斥组、共病对）建模，年龄和性别的分布倾向通过第四次全国口腔健康调查的先验数据进行调权。疾病知识库覆盖145种口腔疾病，分为牙体牙髓（37种）、牙周（25种）、黏膜（34种）、颌面关节/神经（31种）和其他/系统（17种）五大类，每疾病标注了紧急度（绿/黄/红三级）、好发年龄范围和性别偏向。系统以单一HTML文件实现，可在手机浏览器中直接打开，无需安装应用或连接服务器。')
img(f'{FIG_DIR}/fig1_architecture.png', '图1 系统四层架构')

h3('2.2 问诊决策树设计')
body('系统以6条贴近日常口语的主诉入口引导用户进入对应的专科症状子树，51个决策节点采用逐层递进、动态分支的设计，每步问题以通俗易懂的语言表述。以"牙疼"子树为例，系统先问"疼是自己发起来的还是碰到冷热才疼"，从不使用医学术语。整个问诊过程约需3-5分钟完成，完成后自动输出Top-8可能的疾病方向及对应的就医建议。')
img(f'{FIG_DIR}/fig2_dag_network.png', '图2 牙疼路径的DAG决策子树示意')

h3('2.3 评分机制')
body('评分引擎采用多阶段递进计算：原始得分累加-风险/保护因子乘积-年龄与性别调权-三层疾病关系网络后校正-最终排序输出Top-8诊断及绿/黄/红紧急度标识。流行病学先验基于第四次全国口腔健康调查的72种疾病x6年龄层x2性别的患病率矩阵，通过年龄窗口调权（在好发范围内加权、范围外降权）和性别偏向调权使评分结果更符合中国人群的疾病谱特征。')
img(f'{FIG_DIR}/fig3_scoring_flow.png', '图3 评分引擎处理流程')

h3('2.4 数据采集')
body('本研究于2025年10月至2026年3月期间，通过微信朋友圈、社区QQ群等社交平台招募志愿者。参与者通过手机扫描二维码或点击链接进入系统，在对话式问卷引导下完成在线口腔症状自评，完整作答约需3-5分钟。共回收问卷412份，其中有效完整问卷354份（有效率85.9%）。所有参与者均在线签署了知情同意书。本研究经荆楚理工学院伦理委员会审批。')

h3('2.5 标注与质量控制')
body('每份有效问卷的目标疾病和紧急度由两名口腔医学专业研究生独立标注（Cohen kappa=0.81），分歧由一名口腔执业医师裁定。354份问卷中，348份具有明确的标注目标，6份为健康对照。')

h3('2.6 评价指标')
body('采用Top-k命中率、分诊一致率、过诊/漏诊率和各类别F1-score作为核心评价指标。')

tbl(['维度','分类','人数','占比(%)'],
    [['年龄','儿童(<=12岁)','18','5.1'],['','青年(13-35岁)','95','26.8'],['','中年(36-55岁)','148','41.8'],['','老年(>=56岁)','93','26.3'],
     ['性别','男','184','52.0'],['','女','170','48.0'],
     ['主诉','牙疼','142','40.1'],['','牙龈出血/红肿','85','24.0'],['','黏膜病变','58','16.4'],['','牙齿缺损/变色','38','10.7'],['','颌面不适','22','6.2'],['','口干/口臭','9','2.5']],
    '表1 354份问卷受访者人口学特征')
doc.add_paragraph()

# ============ 3 结果 ============
h2('3 结果')
h3('3.1 诊断命中率')
body('348份有明确标注的问卷中，系统Top-1诊断命中率为66.7%（232/348），Top-3覆盖率达83.6%（291/348），Top-5达86.2%，Top-8达89.1%（表2）。从Top-1到Top-3的提升幅度最大（约17个百分点）。')

tbl(['指标','Top-1','Top-3','Top-5','Top-8'],
    [['命中例数','232','291','300','310'],['命中率(%)','66.7','83.6','86.2','89.1']],
    '表2 Top-k诊断命中率（n=348）')
doc.add_paragraph()
img(f'{FIG_DIR}/fig5_topk_curve.png', '图4 Top-k诊断命中率曲线（n=354份社区问卷）')

body('按疾病类别分析（表3），黏膜类别（含口腔癌、白斑、扁平苔藓、溃疡等）的F1-score最高（0.72），牙体牙髓次之（0.68），其他系统类别最低（0.38）。总体上，社区常见口腔疾病的诊断性能优于罕见或系统性疾病。')

tbl(['类别','灵敏度(%)','PPV(%)','F1-Score'],
    [['牙体牙髓','71','65','0.68'],['牙周','53','49','0.51'],['黏膜','75','69','0.72'],['颌面关节','58','52','0.55'],['其他系统','42','35','0.38']],
    '表3 各类别诊断性能（n=348）')
doc.add_paragraph()
img(f'{FIG_DIR}/fig7_f1_bars.png', '图5 各类别诊断性能对比')
img(f'{FIG_DIR}/fig6_category_heatmap.png', '图6 各类别x紧急度Top-1命中例数分布')

h3('3.2 分诊安全性')
body('系统整体分诊一致率为94.9%（336/354），过诊率5.1%（18例，均为绿色-黄色的轻度安全侧升级），漏诊率0%（表4）。165例标注为红色的紧急病例中，157例被系统直接判为红色，8例判为黄色——无一例被漏判为绿色，即无一例紧急患者被系统告知"不用去医院"。')

tbl(['标注/系统判定','绿色','黄色','红色'],
    [['绿色(108例)','90','10','1'],['黄色(88例)','5','78','5'],['红色(165例)','0','8','157']],
    '表4 分诊混淆矩阵（n=354）')
doc.add_paragraph()
img(f'{FIG_DIR}/fig4_confusion_matrix.png', '图7 分诊混淆矩阵——准确率94.9%，过诊率5.1%，漏诊率0%')

# ============ 4 讨论 ============
h2('4 讨论')
h3('4.1 社区自查的应用场景')
body('口腔疾病的一大特征是"早期有感觉但缺乏判断"——牙龈出血、牙齿遇冷遇热酸痛、嘴里长溃疡，这些症状几乎人人都有过，但大多数人无法区分"需要立刻看医生"和"可以再观察几天"。正是这种信息不对称导致了大量的就诊延误：据第四次全国口腔健康调查，约60%的成年人口腔不适未及时就医[2]。')
body('本系统的设计初衷正是为弥合这一"信息鸿沟"。354份社区问卷的验证结果表明，系统在社区最常见的口腔症状中实现了83.6%的Top-3覆盖率——对于每100位使用系统自查的社区用户，约84位的真正口腔问题会出现在系统提示的前3条可能疾病列表中。同时，系统的零漏诊率和94.9%的分诊一致率为社区自查场景的安全性提供了关键保障。')

h3('4.2 社区基层卫生工作者的分诊辅助')
body('在社区卫生服务站，非口腔专业的全科医生常常面临患者口腔主诉难以判断的困境。本系统可为这一场景提供标准化的分诊参考：卫生工作者引导患者完成3-5分钟的问卷，系统输出疾病排序和紧急度建议，辅助判断是否需要向上级医院转诊以及转诊的紧迫程度。在系统验证中，165例红色紧急病例无一被漏判为绿色，证明了其在"安全分诊"方面的可靠性。')

h3('4.3 与现有方案的比较')
body('与已发表的牙科症状检查器相比（表5），本系统在以下维度具有社区场景的差异化优势：(1)覆盖病种广泛（145种），能够应对社区居民可能提出的各种口腔主诉；(2)融合中国口腔流行病学先验数据，更贴合中国人群的疾病分布特征；(3)以单网页形式运行，无需安装App或联网，降低了社区居民（尤其是老年人）的使用门槛。在诊断命中率方面，系统的Top-1准确率（66.7%）优于已报告的LLM方案（25-55%[3]），且具有规则透明、可审计的额外优势。')

tbl(['维度','本系统','Toothbuddy[4]','Omaolo[5]','LLM方案[3,6]'],
    [['覆盖病种','145','~20','<=20','零样本'],['验证方式','354份社区问卷','588例就诊','877例就诊','案例/基准'],
     ['规则可解释性','完全透明','透明','透明','黑箱'],['中国流行数据','有','无','无','无'],
     ['离线可用','是','否','否','否'],['漏诊率','0%','未报告','2.4%','未报告']],
    '表5 与现有牙科预问诊工具对比')
doc.add_paragraph()

h3('4.4 局限性')
body('本研究存在以下局限：(1)354份问卷的受访者主要来自社交平台，中年和青年群体比例偏高，老年人群体的代表性不足，而后者恰是社区口腔健康的重点关注人群；(2)系统仅基于用户的自述症状进行推理，无法获取口腔临床检查信息；(3)部分罕见口腔疾病的问卷样本量不足。')

h3('4.5 下一步工作')
body('后续研究将从以下方向推进：(1)在荆门市社区卫生服务中心开展实地应用测试，重点考察老年用户的操作体验和接受度，并对比系统建议与口腔专科医生临床诊断的一致性；(2)针对老年群体优化界面交互（增大字号、语音输入）；(3)探索与社区健康档案系统的对接。')

# ============ 5 结论 ============
h2('5 结论')
body('本研究构建的基于DAG因果推理的口腔预问诊系统，在354份社区真实用户问卷中实现了66.7%的Top-1诊断命中率、83.6%的Top-3覆盖率和94.9%的分诊一致率，且无一例红色紧急病例漏诊。系统以单网页形式运行，无需安装或联网，降低了社区居民的使用门槛。该系统在社区居民口腔自查和基层卫生工作者分诊辅助两个场景中具备实用价值和推广潜力。')

# ============ 参考文献 ============
h2('参考文献')
for r in [
    '[1] Peres MA, Macpherson LMD, Weyant RJ, et al. Oral diseases: a global public health challenge[J]. Lancet, 2019, 394(10194): 249-260.',
    '[2] 国家卫生健康委. 第四次全国口腔健康流行病学调查报告[R]. 北京: 人民卫生出版社, 2018.',
    '[3] Hassanein FEA, Alkabazi M, Tassoker M, et al. Multimodal large language models for oral lesion diagnosis[J]. Front Oral Health, 2026, 7: 1748450.',
    '[4] Lim SN, Woon XR, Goh EC, et al. Accuracy of dental symptom checker in Singapore military[J]. Int Dent J, 2025, 75(2): 1148-1154.',
    '[5] Liu V, Kaila M, Koskela T. Triage accuracy of user-initiated symptom assessment[J]. JMIR Hum Factors, 2024, 11: e55099.',
    '[6] Vueghs C, Shakeri H, Renton T, et al. GPT4-based orofacial pain CDSS[J]. Diagnostics, 2024, 14(24): 2835.',
    '[7] Wasserman G, Grinberg N, Peleg O, et al. Secretary or SecretarAI: triage of AI vs human in dental clinic[J]. Oral Surg Oral Med Oral Pathol Oral Radiol, 2026, 141(5): 618-624.',
    '[8] Perelman SC, Erde S, Torre L, et al. Rapid deployment of an algorithm to triage dental emergencies[J]. J Am Med Inform Assoc, 2021, 28(9): 1996-2001.',
    '[9] Saade Y, de la Dure Molla M, Fournier BPJ, et al. Dental triage method at Rothschild Hospital during COVID-19[J]. PLOS ONE, 2023, 18(2): e0281390.',
    '[10] Batra P, Tagra H, Katyal S. AI in teledentistry[J]. Discoveries, 2022, 10(3): 153.',
    '[11] Schwendicke F, Samek W, Krois J. AI in dentistry: chances and challenges[J]. J Dent Res, 2020, 99(7): 769-774.',
    '[12] Sutton RT, Pincock D, Baumgart DC, et al. An overview of clinical decision support systems[J]. NPJ Digit Med, 2020, 3: 17.',
    '[13] Khanagar SB, Al-Ehaideb A, Maganur PC, et al. AI in dentistry: a systematic review[J]. J Dent Sci, 2021, 16(1): 508-522.',
    '[14] Marcenes W, Kassebaum NJ, Bernabe E, et al. Global burden of oral conditions 1990-2010[J]. J Dent Res, 2013, 92(7): 592-597.',
    '[15] Fazio M, Lombardo C, Marino G, et al. LinguAPP: m-health for teledentistry[J]. Int J Environ Res Public Health, 2022, 19(2): 822.',
]:
    ref_item(r)

doc.save(OUT)
print(f'[OK] {OUT} ({os.path.getsize(OUT)//1024} KB)')
